import os
import re
from typing import List
import docx
import fitz  # PyMuPDF
from pypdf import PdfReader
from rapidocr_onnxruntime import RapidOCR
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from src.utils import setup_logger

logger = setup_logger('document_processing')

class DocumentProcessor:
    def __init__(self):
        # Initialize RapidOCR
        # It's lighter and doesn't have the dependency hell of PaddleOCR
        self.ocr = RapidOCR()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=150,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
        )

    def clean_text(self, text: str) -> str:
        lines = text.splitlines()
        cleaned_lines = []
        for line in lines:
            raw = line.strip()
            if not raw:
                continue
            lower = raw.lower()
            if re.fullmatch(r'[0-9ivxlcdm\.]+', lower):
                continue
            if re.search(r'page\s*\d+', lower):
                continue
            if re.search(r'https?://', lower):
                continue
            if '@' in lower and ' ' not in lower:
                continue
            if len(raw) <= 3 and all(ch in '-_=·•—~*·. ' for ch in raw):
                continue
            
           
            special_chars = re.findall(r'[^\u4e00-\u9fa5a-zA-Z0-9\s\.\,\，\。\、\(\)（）]', raw)
            if len(raw) > 5 and len(special_chars) / len(raw) > 0.4:
                continue
                
            # Remove isolated single characters or very short meaningless lines
            # But keep short headings like "1. 引言"
            if len(raw) < 4 and not re.match(r'^\d', raw) and not re.match(r'^第', raw):
                # Check if it's likely noise
                if not re.search(r'[\u4e00-\u9fa5]', raw): # No Chinese characters
                    continue

            cleaned_lines.append(raw)
            
        text = ' '.join(cleaned_lines)
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        return text

    def is_heading_line(self, line: str) -> bool:
        if not line:
            return False
        stripped = line.strip()
        if stripped in ["摘要", "结论", "参考文献", "Abstract"]:
            return True
        if stripped.startswith("第") and ("章" in stripped[:10] or "节" in stripped[:10]):
            return True
        if re.match(r'^\d+(\.\d+)*\s+\S+', stripped):
            return True
        return False

    def split_with_headings(self, text: str) -> List[str]:
        lines = text.splitlines()
        sections: List[str] = []
        current: List[str] = []
        for line in lines:
            if self.is_heading_line(line) and current:
                sections.append("\n".join(current))
                current = [line]
            else:
                current.append(line)
        if current:
            sections.append("\n".join(current))
        chunks: List[str] = []
        for sec in sections:
            cleaned = self.clean_text(sec)
            if not cleaned:
                continue
            sec_chunks = self.text_splitter.split_text(cleaned)
            chunks.extend(sec_chunks)
        return chunks

    def process(self, file_path: str) -> List[Document]:
        """
        Main entry point for processing documents
        Returns list of Document objects with metadata
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return []
            
        ext = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        try:
            if ext == '.docx':
                return self.process_docx(file_path, file_name)
            elif ext == '.pdf':
                return self.process_pdf(file_path, file_name)
            elif ext == '.txt':
                return self.process_txt(file_path, file_name)
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return []
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return []

    def process_docx(self, file_path: str, file_name: str) -> List[Document]:
        """
        Extract text from Word document
        """
        try:
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            raw_text = '\n'.join(full_text)
            chunks = self.split_with_headings(raw_text)
            return [Document(page_content=c, metadata={"source": file_name, "page": 1}) for c in chunks]
        except Exception as e:
            logger.error(f"Error processing docx {file_path}: {e}")
            return []

    def process_txt(self, file_path: str, file_name: str) -> List[Document]:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
            chunks = self.split_with_headings(raw_text)
            return [Document(page_content=c, metadata={"source": file_name, "page": 1}) for c in chunks]
        except Exception as e:
            logger.error(f"Error processing txt {file_path}: {e}")
            return []

    def process_pdf(self, file_path: str, file_name: str) -> List[Document]:
        """
        Extract text from PDF (Text-based or Scanned)
        """
        try:
            reader = PdfReader(file_path)
            documents = []
            is_scanned = True
            
            check_text = ""
            for i in range(min(3, len(reader.pages))):
                check_text += reader.pages[i].extract_text() or ""
            
            if len(check_text.strip()) >= 10:
                is_scanned = False
                logger.info(f"PDF {file_name} identified as Text PDF.")
            else:
                logger.info(f"PDF {file_name} identified as Scanned PDF. Using OCR.")
                return self.process_scanned_pdf(file_path, file_name)

            raw_pages = []
            for page in reader.pages:
                text = page.extract_text() or ""
                raw_pages.append(text)

            header_footer_candidates = {}
            page_lines = []
            for text in raw_pages:
                lines = [l.strip() for l in text.splitlines()]
                non_empty = [l for l in lines if l]
                page_lines.append(non_empty)
                candidates = []
                if non_empty:
                    candidates.append(non_empty[0])
                if len(non_empty) >= 2:
                    candidates.append(non_empty[1])
                if len(non_empty) >= 3:
                    candidates.append(non_empty[-1])
                if len(non_empty) >= 4:
                    candidates.append(non_empty[-2])
                for c in candidates:
                    header_footer_candidates[c] = header_footer_candidates.get(c, 0) + 1

            repeated_lines = set()
            min_pages_for_header = 3
            for line, count in header_footer_candidates.items():
                if count >= min_pages_for_header:
                    repeated_lines.add(line)

            for i, lines in enumerate(page_lines):
                filtered = []
                for l in lines:
                    if l in repeated_lines:
                        continue
                    filtered.append(l)
                page_text = '\n'.join(filtered)
                if not page_text:
                    continue
                page_chunks = self.split_with_headings(page_text)
                for chunk in page_chunks:
                    documents.append(Document(
                        page_content=chunk,
                        metadata={"source": file_name, "page": i + 1}
                    ))
            return documents

        except Exception as e:
            logger.error(f"Error processing pdf {file_path}: {e}")
            return []

    def process_scanned_pdf(self, file_path: str, file_name: str) -> List[Document]:
        """
        Use RapidOCR + PyMuPDF to extract text from scanned PDF
        """
        try:
            doc = fitz.open(file_path)
            documents = []
            
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    # Render page to image (zoom=2 for better quality)
                    mat = fitz.Matrix(2, 2)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # RapidOCR accepts bytes directly
                    img_bytes = pix.tobytes("png")
                    
                    result, _ = self.ocr(img_bytes)
                    
                    page_text = ""
                    if result:
                        for line in result:
                            if line and len(line) >= 2:
                                 page_text += line[1] + "\n"
                    
                    if page_text:
                        clean_text = self.clean_text(page_text)
                        if not clean_text:
                            continue
                        page_chunks = self.text_splitter.split_text(clean_text)
                        for chunk in page_chunks:
                            if not chunk.strip():
                                continue
                            documents.append(Document(
                                page_content=chunk, 
                                metadata={"source": file_name, "page": page_num + 1}
                            ))
                except Exception as e:
                    logger.warning(f"Error processing page {page_num+1} of {file_name}: {e}")
                    continue
            
            return documents
            
        except Exception as e:
            logger.error(f"Error processing scanned pdf {file_path}: {e}")
            return []
            
    # Deprecated legacy method
    def get_chunks(self, file_path: str) -> List[str]:
        docs = self.process(file_path)
        return [d.page_content for d in docs]

# Singleton
doc_processor = DocumentProcessor()
